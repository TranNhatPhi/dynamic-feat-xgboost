"""Dựng bản thảo .docx đẹp: title, abstract, các mục, Bảng 1 (bảng Word), Hình 1 & 2 nhúng."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "results"

TITLE = ("Rethinking Embedded Feature Engineering in Gradient Boosting: "
         "A Static Cheap Rotation Rivals Ensemble and Adaptive Selection")

ABSTRACT = (
    "Embedding feature engineering (FE) into the boosting loop—rather than applying it once "
    "as preprocessing—has been shown to improve gradient-boosted trees. A recent method, "
    "Feat-XGBoost, rotates an ensemble of six FE transforms across boosting iterations. This "
    "design raises two unexamined questions: is the full six-transform ensemble necessary, and "
    "would adaptively selecting a transform per iteration (rather than blindly rotating) do "
    "better? We conduct a controlled empirical study on 15 high-dimensional UCI classification "
    "datasets (4-fold cross-validation, 3 seeds), comparing six selection strategies within a "
    "single boosting framework: no FE, one fixed transform, a 2-transform rotation, a cheap "
    "5-transform rotation, the original 6-transform round-robin, and a contextual-bandit "
    "selector. We find that (i) FE diversity helps—round-robin significantly outperforms plain "
    "XGBoost (Δacc = +2.4 pts, paired t = 2.57) and single-transform variants; (ii) the most "
    "expensive transform (Autofeat) is nearly free of benefit—dropping it reduces peak "
    "feature-representation memory by 57% with no significant accuracy change (Δacc = −0.16 pts, "
    "t = −0.56); and (iii) online adaptive selection provides no advantage: a contextual bandit "
    "ties the cheap static rotation on accuracy (t = −0.37) while consuming as much memory as "
    "the full ensemble, because it must materialize every candidate transform. An oracle "
    "analysis explains the negative result: the per-iteration optimal transform is nearly "
    "uniform across boosting phases, so there is little exploitable structure. We recommend a "
    "static rotation over cheap transforms and release all code and data splits."
)
KEYWORDS = ("gradient boosting, feature engineering, XGBoost, contextual bandit, "
            "empirical study, computational efficiency.")

# (heading, [paragraphs])
SECTIONS = [
    ("1. Introduction", [
        "Feature engineering (FE) frequently determines the effectiveness of tabular machine-"
        "learning models: even strong classifiers struggle when the input representation obscures "
        "class structure. In the gradient-boosting family (XGBoost, LightGBM), FE has traditionally "
        "been a preprocessing step applied once, decoupled from the learner. A recent line of work "
        "instead embeds FE inside the boosting loop: at each iteration a feature transform is "
        "applied to (a subsample of) the training data before the next tree is fit. Feat-XGBoost "
        "implements this idea, rotating six transforms—identity, Autofeat, random projection, a "
        "hard-thresholded SVD, robust scaling, and min–max scaling—across iterations, and reports "
        "accuracy gains over standard baselines on UCI benchmarks.",
        "While effective, this design leaves two questions unanswered. First, is the full ensemble "
        "necessary? Rotating six transforms—some expensive (Autofeat expands the feature space "
        "several-fold)—incurs substantial cost, and each transform's contribution is unclear. "
        "Second, is blind rotation optimal? The transform is chosen by a fixed schedule (iteration "
        "index modulo six), ignoring the data state; adaptively choosing per iteration might be "
        "superior.",
        "We answer these with a controlled empirical study rather than by proposing a new state-of-"
        "the-art method, organized around four research questions. RQ1: Does FE diversity in the "
        "boosting loop help, relative to no FE or a single transform? RQ2: Is the full six-transform "
        "ensemble needed, or does a cheaper subset suffice? RQ3: Does adaptive per-iteration "
        "selection (a contextual bandit) outperform fixed selection? RQ4: Does the optimal transform "
        "vary systematically across boosting iterations, and can this be exploited online?",
        "Contributions: (1) a fair, single-framework comparison of six FE-selection strategies over "
        "15 high-dimensional datasets with statistical testing; (2) evidence that the most expensive "
        "transform (Autofeat) can be removed with a 57% reduction in feature-representation memory "
        "and no significant accuracy loss; (3) a negative but well-supported result that online "
        "adaptive selection does not beat a cheap static rotation, with an oracle analysis "
        "explaining why; (4) a concrete practical recommendation and an open, reproducible codebase.",
    ]),
    ("2. Related Work", [
        "[TO EXPAND] Feature engineering for boosting: automatic feature construction (e.g., "
        "Autofeat), dimensionality reduction (random projection, SVD-based denoising). Prior work "
        "largely treats FE as preprocessing; Feat-XGBoost is closest to ours in embedding FE within "
        "boosting.",
        "Gradient-based one-side sampling (GOSS), from LightGBM, keeps large-gradient (hard) samples "
        "and subsamples small-gradient ones with re-weighting; we use it to focus each iteration on "
        "hard examples. Contextual bandits (e.g., LinUCB) select actions from context to maximize "
        "reward; we repurpose one to select an FE transform per boosting iteration.",
    ]),
    ("3. Materials and Methods", [
        "All strategies share one boosting engine differing only in an interchangeable FE selector. "
        "At iteration t: (1) compute gradients of the multiclass log-loss on the training set; "
        "(2) apply GOSS (a = 0.4 large-gradient, b = 0.3 small-gradient, re-weighting (1−a)/b) to "
        "obtain a hard subset; (3) a selector picks one transform from a candidate set; (4) a single "
        "boosting round (multi:softprob) is fit on the transformed data and added to the additive "
        "margin. At inference, each stored tree is applied in its own transformed space. Candidate "
        "transforms follow Feat-XGBoost: identity, Autofeat (first-round only: 1/x, x^2, x^3, e^x), "
        "Gaussian random projection, hard-thresholded SVD (Gavish–Donoho threshold), robust scaling, "
        "min–max scaling.",
        "The contextual bandit uses a low-dimensional summary of the hard subset as context "
        "(mean/std gradient magnitude, class coverage/entropy, dimensionality); its reward is "
        "validation-accuracy improvement minus a cost penalty. A 10-iteration warm-up forces round-"
        "robin so it observes every arm.",
        "Data and protocol: 15 high-dimensional datasets from the standardized UCI benchmark "
        "(excluding one dataset with a train/test label mismatch), each evaluated with 4-fold cross-"
        "validation × 3 seeds; we report mean ± std accuracy. Because our claims are relative "
        "(strategy vs. strategy under identical conditions), we fix boosting hyperparameters "
        "(η = 0.3, max depth = 6, 100 iterations) across all strategies. Cost is measured on a "
        "single machine: wall-clock fit time (mean of 3 repeats) and the memory footprint of the "
        "precomputed feature representations. Each strategy precomputes only the transforms it can "
        "use; the bandit must precompute all six (its action menu), reflected in its measured cost.",
    ]),
    ("4. Results", [
        "__H__4.1 RQ1–RQ2: diversity helps, but the full ensemble is wasteful",
        "Table 1 reports mean accuracy and cost across strategies of increasing FE diversity. "
        "Accuracy rises from plain XGBoost (0.757) to round-robin (0.781). Round-robin significantly "
        "beats plain XGBoost (Δ = +0.024, paired t = 2.57, 12/15 wins) and single-transform "
        "Fixed-HT-SVD (11/15 wins), confirming RQ1: FE diversity is beneficial. However, the gain "
        "saturates. Dropping Autofeat—the most expensive transform—moves from round-robin (6) to the "
        "cheap 5-transform rotation with no significant accuracy change (Δ = −0.0016, t = −0.56) "
        "while reducing precomputed-representation memory from 5.06 MB to 2.19 MB (57% reduction) and "
        "fit time by ~15%. This answers RQ2: the full ensemble is unnecessary.",
        "__T__",  # placeholder for Table 1
        "__H__4.2 RQ3: adaptive selection does not help—and is dominated",
        "The contextual bandit matches the cheap static rotation on accuracy (Δ = −0.0012, "
        "t = −0.37; no significant difference). It exceeds the impoverished 2-FE (t = 2.16) and "
        "single-FE (t = 2.04) baselines only because those are too narrow; it never beats the cheap "
        "rotation. On the accuracy–cost plane (Figure 1) the bandit is Pareto-dominated by the cheap "
        "rotation: it uses the same memory as the full round-robin (5.06 MB, since it must "
        "precompute Autofeat as a candidate) yet is less accurate than the cheaper rotation. This "
        "answers RQ3: online adaptive selection provides no benefit here and is strictly worse in "
        "cost than a fixed cheap rotation.",
        "__F1__",  # Figure 1
        "__H__4.3 RQ4: why adaptivity fails",
        "To test whether any per-iteration selector could help, we compute a greedy oracle that, at "
        "each iteration, tries all six transforms and keeps the one minimizing validation loss. "
        "Aggregated over datasets, the oracle's transform distribution is nearly uniform across "
        "boosting phases (Figure 2), with only a mild monotone trend: Autofeat is modestly favored "
        "early (22%) and declines late (13%), while identity is common throughout (30–36%). The "
        "clean denoise-early/preserve-late pattern holds only for individual noisy datasets, not in "
        "aggregate. With no strong temporal structure and a noisy per-iteration reward, an online "
        "learner has little to exploit—explaining the RQ3 result.",
        "__F2__",  # Figure 2
    ]),
    ("5. Discussion", [
        "Practical recommendation. For embedding FE into gradient boosting, practitioners should use "
        "a static rotation over cheap transforms (dimensionality reducers and scalers), omitting "
        "expensive feature-expansion methods such as Autofeat. This retains the accuracy benefit of "
        "FE diversity at roughly one-third the memory of the full ensemble and requires no learned "
        "controller.",
        "Why the bandit was worth testing—and why it lost. Blind rotation is an obvious target for "
        "improvement, and a cost-aware bandit can concentrate on cheap transforms. But two factors "
        "neutralize it: (a) it must precompute every candidate transform to keep them selectable, "
        "erasing its cost advantage; and (b) the per-iteration signal is too noisy and too weakly "
        "structured for the learned policy to beat a fixed schedule on accuracy. This delimits when "
        "adaptive FE selection is not worth its complexity.",
        "Limitations. Our comparison is relative under fixed hyperparameters; absolute numbers would "
        "shift with per-configuration tuning (e.g., Optuna), which we did not run and leave as future "
        "work. We study 15 high-dimensional datasets and XGBoost only; base-learner generality "
        "(LightGBM, random forests) and larger datasets remain open. Timing was measured on small "
        "datasets where fit time is dominated by fixed overheads, so the runtime gap is modest; the "
        "memory gap is the more robust cost signal.",
    ]),
    ("6. Conclusion", [
        "Embedding feature engineering into gradient boosting helps, but the popular six-transform "
        "ensemble is over-engineered: a cheap static rotation matches it at a fraction of the cost, "
        "and a learned online selector adds complexity without benefit. We recommend the simple "
        "cheap rotation and release code, data splits, and all experimental artifacts for "
        "reproducibility.",
    ]),
]

TABLE1 = [
    ["Strategy", "#FE", "Accuracy", "Fit time (s)", "FE memory (MB)"],
    ["Plain XGBoost", "0", "0.757", "1.04", "1.02"],
    ["Fixed-HT-SVD", "1", "0.767", "1.30", "0.19"],
    ["2-FE rotation", "2", "0.771", "1.27", "0.47"],
    ["Cheap rotation", "5", "0.779", "1.25", "2.19"],
    ["Round-robin (orig.)", "6", "0.781", "1.47", "5.06"],
    ["Contextual bandit", "6*", "0.778", "1.33", "5.06"],
]

REFS = [
    "Chen, T., Guestrin, C. XGBoost: A Scalable Tree Boosting System. KDD 2016. [verify]",
    "Ke, G. et al. LightGBM: A Highly Efficient Gradient Boosting Decision Tree. NeurIPS 2017. [verify]",
    "Li, L. et al. A Contextual-Bandit Approach to Personalized News Article Recommendation. WWW 2010. [verify]",
    "Fernández-Delgado, M. et al. Do we need hundreds of classifiers to solve real world "
    "classification problems? JMLR 2014. [verify]",
    "Horn, F., Pack, R., Rieger, M. The autofeat Python Library for Automated Feature Engineering "
    "and Selection. 2020. [verify]",
    "Gavish, M., Donoho, D. The Optimal Hard Threshold for Singular Values is 4/√3. IEEE TIT 2014. [verify]",
    "Feat-XGBoost (base paper). Pattern Recognition, 2026. DOI:10.1016/j.patcog.2026.113169. "
    "[verify authors/title]",
]


def set_base_style(doc):
    st = doc.styles["Normal"].font
    st.name = "Times New Roman"
    st.size = Pt(11)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table1(doc):
    t = doc.add_table(rows=len(TABLE1), cols=len(TABLE1[0]))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(TABLE1):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = val
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9.5)
                    if i == 0:
                        run.font.bold = True
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("Table 1. Mean accuracy and cost over 15 datasets (4-fold × 3 seeds; fixed "
                    "hyperparameters). *The bandit chooses one transform per iteration but must "
                    "materialize all six.")
    r.font.italic = True; r.font.size = Pt(9)


def add_figure(doc, path, caption):
    if not path.exists():
        add_body(doc, f"[MISSING FIGURE: {path.name}]")
        return
    doc.add_picture(str(path), width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.italic = True; r.font.size = Pt(9)


def add_heading3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.bold = True; r.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(6)


def main():
    doc = Document()
    set_base_style(doc)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE); r.font.bold = True; r.font.size = Pt(15)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Working draft — Empirical Study. Citations marked [verify] must be checked "
                    "before submission.")
    r.font.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x99, 0x33, 0x33)

    # Abstract
    add_heading3(doc, "Abstract")
    ap = add_body(doc, ABSTRACT)
    for run in ap.runs:
        run.font.italic = True
    kp = doc.add_paragraph()
    r = kp.add_run("Keywords: "); r.font.bold = True; r.font.size = Pt(10)
    r2 = kp.add_run(KEYWORDS); r2.font.size = Pt(10)

    # Sections
    for heading, blocks in SECTIONS:
        h = doc.add_paragraph()
        r = h.add_run(heading); r.font.bold = True; r.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(10)
        for b in blocks:
            if b == "__T__":
                add_table1(doc)
            elif b == "__F1__":
                add_figure(doc, OUT / "fig1_pareto.png",
                           "Figure 1. Accuracy vs. FE memory (Pareto frontier). The cheap rotation "
                           "is the knee; the contextual bandit lies below the frontier (dominated).")
            elif b == "__F2__":
                add_figure(doc, OUT / "fig2_oracle_phase.png",
                           "Figure 2. Oracle-optimal FE transform distribution by boosting phase. "
                           "The distribution is near-uniform, indicating little exploitable temporal "
                           "structure (RQ4).")
            elif b.startswith("__H__"):
                add_heading3(doc, b[5:])
            else:
                add_body(doc, b)

    # References
    h = doc.add_paragraph()
    r = h.add_run("References (TO VERIFY — do not submit unchecked)")
    r.font.bold = True; r.font.size = Pt(13)
    for i, ref in enumerate(REFS, 1):
        add_body(doc, f"[{i}] {ref}")

    dest = ROOT / "docs" / "manuscript.docx"
    doc.save(str(dest))
    print(f"Đã lưu {dest}")


if __name__ == "__main__":
    main()
